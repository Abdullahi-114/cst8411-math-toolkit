from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "CST8411-Math-Toolkit-User-Guide.docx"
REPOSITORY_URL = "https://github.com/Abdullahi-114/cst8411-math-toolkit"
RELEASE_URL = f"{REPOSITORY_URL}/releases/tag/v1.0.0"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 98, 108)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_properties)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shading)
    run = paragraph.add_run(code)
    set_run_font(run, name="Consolas", size=9, color=NAVY)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(section):
    header = section.header
    header.distance = Inches(0.492)
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("CST8411 Math Toolkit  |  User Guide")
    set_run_font(run, size=9, color=GRAY, bold=True)

    footer = section.footer
    footer.distance = Inches(0.492)
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Assignment 1 - Deployable JAR  |  Version 1.0.0")
    set_run_font(run, size=8.5, color=GRAY)


def add_metadata_table(doc):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    table.cell(0, 0).text = "Release detail"
    table.cell(0, 1).text = "Value"
    mark_header_row(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = NAVY
    rows = [
        ("Library", "CST8411 Math Toolkit"),
        ("Artifact", "math-toolkit-1.0.0.jar"),
        ("Java", "JRE/JDK 17 or later"),
        ("Dependencies", "None"),
    ]
    for index, (label, value) in enumerate(rows, start=1):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
        set_cell_shading(table.cell(index, 0), LIGHT_BLUE)
        for run in table.cell(index, 0).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()


def add_api_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1800, 3900, 3660])
    headers = ["Method", "Purpose", "Example result"]
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = NAVY
    mark_header_row(table.rows[0])

    rows = [
        ("add", "Adds two values", "add(12, 8) = 20"),
        ("subtract", "Subtracts the second value", "subtract(12, 8) = 4"),
        ("multiply", "Multiplies two values", "multiply(6, 7) = 42"),
        ("divide", "Divides dividend by divisor", "divide(20, 4) = 5"),
        ("modulus", "Returns the division remainder", "modulus(17, 5) = 2"),
        ("power", "Raises a base to an exponent", "power(2, 3) = 8"),
        ("squareRoot", "Returns the principal square root", "squareRoot(81) = 9"),
        ("percentage", "Returns part as percent of whole", "percentage(25, 200) = 12.5"),
    ]
    for method, purpose, example in rows:
        cells = table.add_row().cells
        cells[0].text = method
        cells[1].text = purpose
        cells[2].text = example
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, [1800, 3900, 3660])


def build_document():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    add_header_footer(section)

    # Editorial cover for the library's user-facing guide.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(64)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("JAVA LIBRARY USER GUIDE")
    set_run_font(run, size=10.5, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("CST8411 Math Toolkit")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A reusable, dependency-free arithmetic library for Java 17+")
    set_run_font(run, size=14, color=DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(26)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Assignment 1 - Deployable JAR  |  Version 1.0.0  |  June 2026")
    set_run_font(run, size=10.5, color=GRAY, bold=True)
    meta.paragraph_format.space_after = Pt(88)

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prepared.add_run("Prepared by Mohamed Abdullahi")
    set_run_font(run, size=11, color=NAVY, bold=True)

    doc.add_page_break()

    doc.add_heading("1. Library Overview", level=1)
    doc.add_paragraph(
        "CST8411 Math Toolkit packages common arithmetic operations behind a small, reusable Java API. "
        "It can be used as a normal library dependency or run directly as a command-line application. "
        "The release is self-contained and has no third-party runtime dependencies."
    )
    add_metadata_table(doc)

    doc.add_heading("Key capabilities", level=2)
    for item in (
        "Reusable MathOperations interface and StandardMathOperations implementation.",
        "Static MathToolkit facade for concise calls from consuming applications.",
        "Executable JAR with a manifest-defined command-line entry point.",
        "Validation for division by zero, modulus by zero, invalid square roots, and invalid percentages.",
        "Java 17 bytecode compatibility and dependency-free runtime distribution.",
    ):
        add_bullet(doc, item)

    doc.add_heading("2. Prerequisites", level=1)
    add_bullet(doc, "JRE 17 or later to run the packaged JAR.")
    add_bullet(doc, "JDK 17 or later to compile the project or integrate it manually.")
    add_bullet(doc, "PowerShell is used by the included build.ps1 automation script.")
    add_bullet(doc, "No external Java libraries are required at runtime.")

    doc.add_heading("3. Download and Verify", level=1)
    doc.add_paragraph("The source code and release artifacts are published in the project repository:")
    paragraph = doc.add_paragraph()
    add_hyperlink(paragraph, "Open the CST8411 Math Toolkit repository", REPOSITORY_URL)
    doc.add_paragraph("Version 1.0.0 release page:")
    paragraph = doc.add_paragraph()
    add_hyperlink(paragraph, "Open the version 1.0.0 release", RELEASE_URL)

    doc.add_heading("Verify the downloaded file", level=2)
    doc.add_paragraph(
        "The release includes SHA256SUMS.txt. Compare the downloaded JAR checksum with that file before use."
    )
    add_code_block(
        doc,
        "Get-FileHash -Algorithm SHA256 .\\math-toolkit-1.0.0.jar",
    )

    doc.add_heading("4. Add the Library as a Dependency", level=1)
    doc.add_heading("Option A - Direct JAR dependency", level=2)
    for step in (
        "Create a lib directory in the consuming Java project.",
        "Copy math-toolkit-1.0.0.jar into the lib directory.",
        "Add the JAR to the compile and runtime classpath.",
    ):
        add_numbered(doc, step)
    add_code_block(
        doc,
        'javac -cp "lib\\math-toolkit-1.0.0.jar" MyApplication.java\n'
        'java -cp ".;lib\\math-toolkit-1.0.0.jar" MyApplication',
    )

    doc.add_heading("Option B - Install into a local Maven repository", level=2)
    doc.add_paragraph(
        "Developers with Maven installed can register the downloaded JAR locally, then reference it with normal dependency metadata."
    )
    add_code_block(
        doc,
        "mvn install:install-file `\n"
        "  -Dfile=math-toolkit-1.0.0.jar `\n"
        "  -DgroupId=ca.algonquin.cst8411 `\n"
        "  -DartifactId=math-toolkit `\n"
        "  -Dversion=1.0.0 `\n"
        "  -Dpackaging=jar",
    )
    add_code_block(
        doc,
        "<dependency>\n"
        "    <groupId>ca.algonquin.cst8411</groupId>\n"
        "    <artifactId>math-toolkit</artifactId>\n"
        "    <version>1.0.0</version>\n"
        "</dependency>",
    )

    doc.add_heading("5. Java Usage Examples", level=1)
    doc.add_paragraph("Import MathToolkit and call the required operation from application code:")
    add_code_block(
        doc,
        "import ca.algonquin.cst8411.mathtoolkit.MathToolkit;\n\n"
        "public class MyApplication {\n"
        "    public static void main(String[] args) {\n"
        "        double total = MathToolkit.add(15, 27);\n"
        "        double remainder = MathToolkit.modulus(17, 5);\n"
        "        double percent = MathToolkit.percentage(45, 60);\n\n"
        "        System.out.println(total);      // 42.0\n"
        "        System.out.println(remainder);  // 2.0\n"
        "        System.out.println(percent);    // 75.0\n"
        "    }\n"
        "}",
    )

    doc.add_heading("Use the interface for dependency injection", level=2)
    add_code_block(
        doc,
        "import ca.algonquin.cst8411.mathtoolkit.MathOperations;\n"
        "import ca.algonquin.cst8411.mathtoolkit.StandardMathOperations;\n\n"
        "MathOperations operations = new StandardMathOperations();\n"
        "double result = operations.power(2, 10); // 1024.0",
    )

    doc.add_heading("6. Command-Line Usage", level=1)
    doc.add_paragraph("The manifest identifies MathToolkitDemo as the entry point, so the JAR can run directly:")
    add_code_block(
        doc,
        "java -jar math-toolkit-1.0.0.jar add 12 8\n"
        "java -jar math-toolkit-1.0.0.jar sqrt 81\n"
        "java -jar math-toolkit-1.0.0.jar percentage 25 200",
    )
    doc.add_paragraph("Expected output:")
    add_code_block(doc, "Result: 20.0\nResult: 9.0\nResult: 12.5")

    doc.add_heading("7. API Reference", level=1)
    add_api_table(doc)

    doc.add_heading("8. Error Handling", level=1)
    doc.add_paragraph(
        "The toolkit throws ArithmeticException when an operation is mathematically undefined in the supported real-number domain."
    )
    add_bullet(doc, "divide and modulus reject a zero divisor.")
    add_bullet(doc, "squareRoot rejects negative input.")
    add_bullet(doc, "percentage rejects a zero whole value.")
    add_code_block(
        doc,
        "try {\n"
        "    MathToolkit.divide(10, 0);\n"
        "} catch (ArithmeticException exception) {\n"
        "    System.out.println(exception.getMessage());\n"
        "}",
    )

    doc.add_heading("9. Build and Test from Source", level=1)
    for step in (
        "Clone or download the repository.",
        "Open PowerShell in the project root.",
        "Run .\\build.ps1.",
        "Confirm that all 12 checks pass and the dist directory is created.",
    ):
        add_numbered(doc, step)
    add_code_block(doc, ".\\build.ps1")
    doc.add_paragraph(
        "The script compiles with --release 17, packages the executable and source JARs, runs the dependency-free tests, and writes SHA-256 checksums."
    )

    doc.add_heading("10. Versioning and Maintenance", level=1)
    doc.add_paragraph(
        "The project uses semantic versioning. Version 1.0.0 is the initial stable release. Future bug fixes should increment the patch number, backward-compatible features should increment the minor number, and breaking API changes should increment the major number."
    )
    add_bullet(doc, "Report problems through the repository Issues page.")
    add_bullet(doc, "Review release notes before upgrading to a new major version.")
    add_bullet(doc, "Verify checksums whenever release artifacts are downloaded.")

    doc.add_heading("11. Submission Checklist", level=1)
    for item in (
        "Deployable executable JAR is present in dist.",
        "Manifest contains the Main-Class entry.",
        "All 12 functional checks pass.",
        "Source and usage documentation are included.",
        "Repository and version 1.0.0 release links are accessible.",
    ):
        add_bullet(doc, item)

    doc.core_properties.title = "CST8411 Math Toolkit User Guide"
    doc.core_properties.subject = "Assignment 1 - Deployable JAR"
    doc.core_properties.author = "Mohamed Abdullahi"
    doc.core_properties.keywords = "Java, JAR, library, CST8411, documentation"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
